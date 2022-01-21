# -*- coding=utf-8 -*-

import os
import datetime

from docx import Document
from docxcompose.composer import Composer


# 合并Word文档
def merge_doc(source_file_path_list, target_file_path):
    '''
    合并多个docx文件
    :param source_file_path_list: 源文件路径列表
    :param target_file_path: 目标文件路径
    '''

    # 填充分页符号文档
    page_break_doc = Document()
    # page_break_doc.add_page_break()

    # 填充换行符号文档
    page_break_doc.add_paragraph('\n\n')

    # 定义新文档
    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(source_file_path_list)):
        # 跳过第一个
        if i == 0:
            continue
        # 填充分页符文档
        target_composer.append(page_break_doc)
        # 拼接文档内容
        f = source_file_path_list[i]
        target_composer.append(Document(f))
    # 保存目标文档
    target_composer.save(target_file_path)


if __name__ == '__main__':
    # 源文件夹路径
    source_path = r'你自己的文件夹路径'
    # 目标文件路径
    target_file = r'你自己的文件夹路径\target.docx'
    source_file_list = os.listdir(source_path)
    # 获取源文件夹内文件列表
    # 这里word文件命名格式为：yyyy-mm-dd例会议题.docx
    # 时间可以是yyyy-m-d等不标准的格式，但必须包含年月日，且用'-'连接
    source_file_list_all = []
    for f in source_file_list:
        if f.find('例会议题') != -1:
            source_file_list_all.append(source_path + '\\' + f)
            print(f)

    print('---------------- 给文件排个序：字典序==>按时间顺序 -------------------')

    file_list = dict()
    for file in source_file_list_all:
        # 格式化时间
        date = file.split("\\")[-1].split('例')[0]
        date = datetime.datetime.strptime(date, "%Y-%m-%d")
        file_list.update({date: file})

    # 存储路径文件列表
    res = []
    for i in sorted(file_list.items()):
        res.append(i[1])
        print(i[1].split('\\')[-1])

    merge_doc(res, target_file)